import shutil
from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import json
import io
import numpy as np
from PIL import Image
import requests
from urllib.parse import urlparse, quote
import hashlib
from llm import chat, get_session_history, clear_session
from langchain_core.messages import HumanMessage, AIMessage
from dotenv import load_dotenv
from history import ChatMessageHistoryWithTimestamps
from chroma import add_single_document, rebuild_database, remove_documents_by_source
from embeddings import embeddings

try:
	import pypdf
	PDF_AVAILABLE = True
except ImportError:
	PDF_AVAILABLE = False
	print("Warning: pypdf not installed. PDF processing will be limited.")
 
 # Check for docx availability
try:
	import docx
	DOCX_AVAILABLE = True
except ImportError:
	DOCX_AVAILABLE = False

# Initialize EasyOCR reader globally
ocr_reader = None
EASYOCR_AVAILABLE = False

def initialize_easyocr():
	"""Initialize EasyOCR reader with error handling."""
	global ocr_reader, EASYOCR_AVAILABLE
	try:
		import easyocr
		ocr_reader = easyocr.Reader(['en'], gpu=False)
		EASYOCR_AVAILABLE = True
		print("EasyOCR initialized successfully")
	except ImportError:
		EASYOCR_AVAILABLE = False
		ocr_reader = None
		print("Warning: easyocr not installed. Image OCR will be limited.")
	except Exception as e:
		EASYOCR_AVAILABLE = False
		ocr_reader = None
		print(f"Warning: easyocr initialization failed: {e}")

# Initialize on startup
initialize_easyocr()
load_dotenv()

DATA_PATH = os.getenv('DATA_PATH', 'data')
TEMP_PATH = os.getenv('TEMP_PATH', 'temp')
# clear and recreate data folder
if os.path.exists(DATA_PATH):
	shutil.rmtree(DATA_PATH)
os.makedirs(DATA_PATH)

from db import get_session_db, clear_session_db  # Import session-aware functions

PORT = int(os.getenv('BACKEND_PORT', 5050))

app = Flask(__name__)
app.config['CORS_HEADERS'] = 'Content-Type'
app.config['CORS_RESOURCES'] = {r"/*": {"origins": "*"}}
app.config['CORS_SUPPORTS_CREDENTIALS'] = True

CORS(app)

def extract_image_text(image_file):
	"""Extract text from image file using EasyOCR."""
	if not EASYOCR_AVAILABLE or ocr_reader is None:
		return f"[Image OCR unavailable - easyocr not installed]"
	
	try:
		# Read image file into memory
		image_data = image_file.read()
		image_file.seek(0)
		
		# Convert bytes to PIL Image
		image = Image.open(io.BytesIO(image_data))
		
		# Convert to RGB if necessary (handles RGBA, grayscale, etc.)
		if image.mode != 'RGB':
			image = image.convert('RGB')
		
		# Convert to numpy array for easyocr
		image_array = np.array(image)
		
		# Run OCR with confidence threshold
		ocr_result = ocr_reader.readtext(image_array, detail=1)
		
		# Extract text from results (ocr_result is list of tuples: (bbox, text, confidence))
		if not ocr_result:
			return f"[Image File: {image_file.filename}]\nNo text detected in image."
		
		# Extract text with confidence filtering
		text_lines = [item[1] for item in ocr_result if item[2] > 0.3]  # confidence > 0.3
		text = '\n'.join(text_lines)
		
		if not text.strip():
			return f"[Image File: {image_file.filename}]\nNo text detected in image."
		
		return text
	except Exception as e:
		print(f"Image OCR Error for {image_file.filename}: {str(e)}")
		return f"[Image OCR Error: {str(e)}]"

def extract_pdf_text(pdf_file):
	"""Extract text from PDF file using pypdf and EasyOCR fallback."""
	if not PDF_AVAILABLE:
		return f"[PDF processing unavailable - pypdf not installed]"
	
	try:
		# Save file position to restore later
		original_pos = pdf_file.stream.tell() if hasattr(pdf_file, 'stream') else 0
		
		pdf_reader = pypdf.PdfReader(pdf_file)
		text = ""
		
		for page_num, page in enumerate(pdf_reader.pages):
			try:
				# Try text extraction first
				page_text = page.extract_text()
				
				if page_text and page_text.strip():
					text += f"\n--- Page {page_num + 1} ---\n"
					text += page_text
				else:
					# If no text extracted, try EasyOCR fallback on the page
					if EASYOCR_AVAILABLE and ocr_reader is not None:
						text += f"\n--- Page {page_num + 1} (OCR) ---\n"
						try:
							# Try to render PDF page to image for OCR
							try:
								import fitz  # pymupdf
								pdf_file.seek(0)
								doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
								page_image = doc[page_num].get_pixmap(matrix=fitz.Matrix(2, 2))
								
								# Convert pixmap to numpy array
								image_data = page_image.tobytes("ppm")
								image = Image.open(io.BytesIO(image_data))
								image_array = np.array(image)
								
								# Run OCR on the page image
								ocr_result = ocr_reader.readtext(image_array, detail=1)
								ocr_text = '\n'.join([item[1] for item in ocr_result if item[2] > 0.3])
								text += ocr_text if ocr_text.strip() else "[No text detected on page]\n"
								doc.close()
							except ImportError:
								text += "[pymupdf not installed - cannot OCR scanned PDF pages]\n"
						except Exception as ocr_error:
							text += f"[OCR failed for page {page_num + 1}: {str(ocr_error)}]\n"
					else:
						text += f"[No text on page {page_num + 1}]\n"
			except Exception as e:
				text += f"[Failed to extract page {page_num + 1}: {str(e)}]\n"
		
		return text if text.strip() else "[PDF file had no extractable text]"
	except Exception as e:
		print(f"PDF processing error: {str(e)}")
		return f"[Failed to process PDF: {str(e)}]"

def extract_docx_text(doc_file):
	"""Extract text from DOCX/DOC file."""
	if not DOCX_AVAILABLE:
		return f"[DOCX processing unavailable - python-docx not installed]"
	
	try:
		# Get filename
		filename = getattr(doc_file, 'filename', 'document')
		file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
		
		# Handle old .doc format (binary)
		if file_extension == 'doc':
			return handle_old_doc_format(doc_file, filename)
		
		# Handle .docx format (modern XML-based)
		# Read file into memory
		file_data = doc_file.read()
		doc_file.seek(0)  # Reset file position
		
		# Load document from bytes
		doc = docx.Document(io.BytesIO(file_data))
		
		# Extract all text from paragraphs
		text_parts = []
		
		# Extract paragraphs
		for paragraph in doc.paragraphs:
			if paragraph.text.strip():
				text_parts.append(paragraph.text)
		
		# Extract text from tables
		for table in doc.tables:
			for row in table.rows:
				row_text = []
				for cell in row.cells:
					if cell.text.strip():
						row_text.append(cell.text.strip())
				if row_text:
					text_parts.append(' | '.join(row_text))
		
		# Extract text from headers and footers
		for section in doc.sections:
			# Header
			if section.header:
				for paragraph in section.header.paragraphs:
					if paragraph.text.strip():
						text_parts.append(f"[Header: {paragraph.text}]")
			
			# Footer
			if section.footer:
				for paragraph in section.footer.paragraphs:
					if paragraph.text.strip():
						text_parts.append(f"[Footer: {paragraph.text}]")
		
		# Combine all text
		text = '\n'.join(text_parts)
		
		if not text.strip():
			return f"[Document File: {filename}]\nNo text detected in document."
		
		return text
		
	except Exception as e:
		print(f"DOCX processing error for {filename}: {str(e)}")
		return f"[Failed to process document: {str(e)}]"


def handle_old_doc_format(doc_file, filename):
	"""Handle old .doc format (binary/OLE format)."""
	# Old .doc files are complex binary format
	# Best approach: inform user to convert to .docx or use OCR on PDF export
	
	return f"[Old .doc format detected for {filename}. Please convert to .docx or PDF for better text extraction.]"


def extract_url_text(url: str) -> str:
	"""Extract text from a URL by taking a screenshot and running OCR."""
	print(f"Extracting text from URL: {url}")
	
	# Try Selenium + OCR first (best quality for visual content)
	screenshot_image = None
	try:
		from selenium import webdriver
		from selenium.webdriver.chrome.options import Options
		
		print("Attempting Selenium screenshot...")
		# Setup Chrome options
		chrome_options = Options()
		chrome_options.add_argument('--headless')
		chrome_options.add_argument('--no-sandbox')
		chrome_options.add_argument('--disable-dev-shm-usage')
		chrome_options.add_argument('--disable-gpu')
		chrome_options.add_argument('--window-size=1920,1080')
		
		# Create driver
		driver = webdriver.Chrome(options=chrome_options)
		driver.get(url)
		
		# Wait for page to load
		import time
		time.sleep(3)  # Give page time to render
		
		# Take screenshot
		screenshot_bytes = driver.get_screenshot_as_png()
		driver.quit()
		
		# Convert to PIL Image
		screenshot_image = Image.open(io.BytesIO(screenshot_bytes))
		print("Screenshot captured successfully")
		
	except ImportError:
		print("Selenium not available - will try BeautifulSoup fallback")
	except Exception as selenium_error:
		print(f"Selenium screenshot failed: {selenium_error}")
	
	# If we have a screenshot and OCR is available, use it
	if screenshot_image and EASYOCR_AVAILABLE and ocr_reader is not None:
		print("Running OCR on screenshot...")
		try:
			# Convert to RGB if necessary
			if screenshot_image.mode != 'RGB':
				screenshot_image = screenshot_image.convert('RGB')
			
			# Convert to numpy array for easyocr
			image_array = np.array(screenshot_image)
			
			# Run OCR
			ocr_result = ocr_reader.readtext(image_array, detail=1)
			
			if ocr_result:
				# Extract text with confidence filtering
				text_lines = [item[1] for item in ocr_result if item[2] > 0.3]
				text = '\n'.join(text_lines)
				
				if text.strip():
					print(f"OCR extracted {len(text_lines)} lines of text")
					return f"[Website Content from {url}]\n(Extracted via Screenshot OCR)\n\n{text}"
			
			print("OCR found no text in screenshot")
		except Exception as ocr_error:
			print(f"OCR processing failed: {ocr_error}")
	
	# Fallback: Try BeautifulSoup HTML parsing
	print("Attempting BeautifulSoup text extraction...")
	try:
		from bs4 import BeautifulSoup
		response = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
		response.raise_for_status()
		
		soup = BeautifulSoup(response.content, 'html.parser')
		
		# Remove script and style elements
		for script in soup(["script", "style"]):
			script.decompose()
		
		# Get text
		text = soup.get_text(separator='\n')
		
		# Clean up whitespace
		lines = (line.strip() for line in text.splitlines())
		chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
		text = '\n'.join(chunk for chunk in chunks if chunk)
		
		if text.strip():
			print(f"BeautifulSoup extracted text successfully")
			return f"[Website Content from {url}]\n(Extracted via HTML parsing)\n\n{text}"
		else:
			return f"[Website: {url}]\nNo text content extracted."
			
	except ImportError:
		return f"[Website: {url}]\nText extraction requires BeautifulSoup (bs4) to be installed."
	except Exception as bs_error:
		print(f"BeautifulSoup extraction failed: {bs_error}")
		return f"[Error fetching URL: {url}]\n{str(bs_error)}"

def extract_document_text(doc_file):
	"""
	Unified function to extract text from various document formats.
	Automatically detects format based on file extension.
	"""
	filename = getattr(doc_file, 'filename', 'document')
	file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
	
	if file_extension in ['docx', 'doc']:
		return extract_docx_text(doc_file)
	elif file_extension == 'pdf':
		return extract_pdf_text(doc_file)
	elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
		return extract_image_text(doc_file)
	else:
		return f"[Unsupported file format: {file_extension}]"

@app.route('/')
def home():
	return jsonify({'message': 'Welcome to the Chat API!'})

@app.route('/api/chat', methods=['POST'])
def chat_endpoint():
	session_id = request.headers.get('Session-ID')

	# make session folder if not exists
	session_folder = os.path.join(DATA_PATH, session_id)
	if not os.path.exists(session_folder):
		os.makedirs(session_folder)

	# Handle both JSON and FormData requests
	message = None
	files = []
	file_contents = {}
	urls = []
	
	if request.is_json:
		# Handle JSON request
		data = request.get_json()
		message = data.get('message')
		selected_sources = data.get('selectedSources', [])
	else:
		# Handle FormData request (for file uploads)
		message = request.form.get('message')
		
		# Parse JSON fields from form data
		selected_sources_str = request.form.get('selectedSources', '[]')
		try: 
			selected_sources = json.loads(selected_sources_str)
		except json.JSONDecodeError:
			selected_sources = []

	print("Message:", message)
	print("Selected Sources:", selected_sources)

	if not message:
		return jsonify({'error': 'Input is required'}), 400
	if not session_id or session_id.strip() == "":
		return jsonify({'error': 'Session-ID header is required'}), 400

	print("chat request for session:", session_id)

	# Get session-specific database
	session_db = get_session_db(session_id)

	# Pass files and file_contents to chat function
	# Use os.path.join to ensure correct path separators for the OS
	source_names = [os.path.join(DATA_PATH, session_id, f"{src['name']}.md") for src in selected_sources if src]
	sources = [{'name': source_names[idx], 'size': selected_sources[idx]['size']} for idx in range(len(selected_sources)) if selected_sources[idx]]
	print(f"Selected sources from frontend: {selected_sources}")
	print(f"Constructed source paths: {source_names}")
	
	# Verify that the source files actually exist
	for source_path in source_names:
		if os.path.exists(source_path):
			print(f"Found: {source_path}")
		else:
			print(f"Missing: {source_path}")
	
	reply = chat(message, session_id=session_id, db=session_db, selected_sources=sources)
	return jsonify({'reply': reply})

# upload/delete/get sources
@app.route('/api/sources', methods=['POST', 'DELETE', 'GET'])
def source_endpoint():
	session_id = request.headers.get('Session-ID')
	if not session_id or session_id.strip() == "":
		return jsonify({'error': 'Session-ID header is required'}), 400
	
	if request.method == 'DELETE':
		# delete file by filename
		filename = request.args.get('filename')
		if not filename:
			return jsonify({'error': 'filename parameter is required for DELETE'}), 400
		session_folder = os.path.join(DATA_PATH, session_id)
		file_path = os.path.join(session_folder, filename + '.md')
		if os.path.exists(file_path):
			os.remove(file_path)
			# Remove from session's chroma vector store
			session_db = get_session_db(session_id)
			remove_documents_by_source(session_db, file_path)
			return jsonify({'message': f'File {filename} deleted from session {session_id}'}), 200
		else:
			return jsonify({'error': f'File {filename} not found in session {session_id}'}), 404
	elif request.method == 'POST':
		if 'file' not in request.files:
			return jsonify({'error': 'No file part in the request'}), 400

		file = request.files['file']
		if file.filename == '':
			return jsonify({'error': 'No selected file'}), 400
		
		# use ocr to extract text if image, pdf, or docx
		extracted_text = extract_document_text(file)
		print(f"Extracted text from {file.filename}:\n{extracted_text[:500]}...")  # Print first 500 chars

		# store file in data folder in session subfolder as .md file with extracted text
		session_folder = os.path.join(DATA_PATH, session_id)
		if not os.path.exists(session_folder):
			os.makedirs(session_folder)
		file_path = os.path.join(session_folder, f"{file.filename}.md")
		with open(file_path, 'w', encoding='utf-8') as f:
			f.write(extracted_text)
   
		# save to chroma vector store as well
		session_db = get_session_db(session_id)
		add_single_document(session_db, embeddings, filepath=file_path, session_id=session_id)
		return jsonify({'message': f'File saved to {file_path}'}), 200
	elif request.method == 'GET':
		# list files in session folder (excluding .web.md files - those are listed via /api/urls)
		session_folder = os.path.join(DATA_PATH, session_id)
		if not os.path.exists(session_folder):
			return jsonify({'files': []}), 200
		# Only include .md files that are NOT .web.md files
		file_names = [f[:-3] for f in os.listdir(session_folder) if f.endswith('.md') and not f.endswith('.web.md')]
		extensions = [file.rsplit('.', 1)[-1].lower() for file in file_names]
		files = [{'name': name, 'extension': ext} for name, ext in zip(file_names, extensions)]
		return jsonify({'files': files}), 200

# URL management endpoint
@app.route('/api/urls', methods=['POST', 'DELETE', 'GET'])
def url_endpoint():
	print(f"========== /api/urls {request.method} request received ==========")
	session_id = request.headers.get('Session-ID')
	
	# Allow empty session for POST - we'll create one automatically
	if request.method == 'POST' and (not session_id or session_id.strip() == ""):
		# Generate a new session ID
		import time
		session_id = f"session-{int(time.time() * 1000)}"
		print(f"No session provided, created new session: {session_id}")
	elif not session_id or session_id.strip() == "":
		return jsonify({'error': 'Session-ID header is required'}), 400
	
	print(f"Session ID: {session_id}")
	
	# Create session folder if it doesn't exist
	session_folder = os.path.join(DATA_PATH, session_id)
	if not os.path.exists(session_folder):
		os.makedirs(session_folder)
		print(f"Created session folder: {session_folder}")
	
	if request.method == 'DELETE':
		# Delete URL by filename
		url_hash = request.args.get('urlHash')
		if not url_hash:
			return jsonify({'error': 'urlHash parameter is required for DELETE'}), 400
		
		file_path = os.path.join(session_folder, f"{url_hash}.web.md")
		if os.path.exists(file_path):
			os.remove(file_path)
			# Remove from session's chroma vector store
			session_db = get_session_db(session_id)
			remove_documents_by_source(session_db, file_path)
			return jsonify({'message': f'URL deleted from session {session_id}'}), 200
		else:
			return jsonify({'error': f'URL not found in session {session_id}'}), 404
			
	elif request.method == 'POST':
		try:
			data = request.get_json()
			print(f"Received POST data: {data}")
			url = data.get('url')
			
			if not url:
				print("ERROR: No URL provided in request")
				return jsonify({'error': 'URL is required'}), 400
			
			# Validate URL
			try:
				parsed = urlparse(url)
				if not parsed.scheme or not parsed.netloc:
					return jsonify({'error': 'Invalid URL format'}), 400
			except Exception:
				return jsonify({'error': 'Invalid URL format'}), 400
			
			print(f"Processing URL: {url}")
			
			# Extract text from URL
			extracted_text = extract_url_text(url)
			print(f"Extracted text from {url}:\n{extracted_text[:500]}...")  # Print first 500 chars
			
			# Create a hash of the URL for the filename
			url_hash = hashlib.md5(url.encode()).hexdigest()[:12]
			
			# Store as .web.md file
			file_path = os.path.join(session_folder, f"{url_hash}.web.md")
			print(f"Saving URL content to: {file_path}")
			with open(file_path, 'w', encoding='utf-8') as f:
				f.write(f"URL: {url}\n\n{extracted_text}")
			
			print(f"File saved, size: {os.path.getsize(file_path)} bytes")
			
			# Add to chroma vector store
			session_db = get_session_db(session_id)
			print(f"Adding to vector store for session: {session_id}")
			add_single_document(session_db, embeddings, filepath=file_path, session_id=session_id)
			
			print(f"URL processing complete: {url_hash}.web")
			return jsonify({
				'message': 'URL processed successfully',
				'urlHash': url_hash,
				'url': url,
				'name': f"{url_hash}.web",
				'sessionId': session_id  # Return session ID in case it was auto-created
			}), 200
			
		except Exception as e:
			print(f"ERROR in URL POST: {str(e)}")
			import traceback
			traceback.print_exc()
			return jsonify({'error': f'Failed to process URL: {str(e)}'}), 500
		
	elif request.method == 'GET':
		# List all URLs in session
		print(f"GET URLs - checking folder: {session_folder}")
		if not os.path.exists(session_folder):
			print("Session folder doesn't exist, returning empty list")
			return jsonify({'urls': []}), 200
		
		urls = []
		all_files = os.listdir(session_folder)
		print(f"Files in session folder: {all_files}")
		
		for filename in all_files:
			if filename.endswith('.web.md'):
				print(f"Found URL file: {filename}")
				file_path = os.path.join(session_folder, filename)
				try:
					# Read the first line to get the URL
					with open(file_path, 'r', encoding='utf-8') as f:
						first_line = f.readline().strip()
						if first_line.startswith('URL: '):
							original_url = first_line[5:]  # Remove 'URL: ' prefix
							url_hash = filename[:-7]  # Remove '.web.md'
							url_obj = {
								'url': original_url,
								'urlHash': url_hash,
								'name': f"{url_hash}.web"
							}
							print(f"  Returning URL object: {url_obj}")
							urls.append(url_obj)
				except Exception as e:
					print(f"Error reading URL file {filename}: {e}")
					continue
		
		print(f"Returning {len(urls)} URLs")
		return jsonify({'urls': urls}), 200
	
	elif request.method == 'DELETE':
		# Delete URL by urlHash
		url_hash = request.args.get('urlHash')
		if not url_hash:
			return jsonify({'error': 'urlHash parameter is required'}), 400
		
		print(f"DELETE URL - urlHash: {url_hash}")
		file_path = os.path.join(session_folder, f"{url_hash}.web.md")
		
		if not os.path.exists(file_path):
			print(f"URL file not found: {file_path}")
			return jsonify({'error': 'URL not found'}), 404
		
		try:
			os.remove(file_path)
			print(f"Deleted URL file: {file_path}")
			return jsonify({'message': 'URL deleted successfully'}), 200
		except Exception as e:
			print(f"Error deleting URL file: {e}")
			return jsonify({'error': f'Failed to delete URL: {str(e)}'}), 500

# get all sessions endpoint
@app.route('/api/sessions', methods=['GET'])
def sessions_endpoint():
	sessions = []
	if os.path.exists(DATA_PATH):
		sessions = [name for name in os.listdir(DATA_PATH) if os.path.isdir(os.path.join(DATA_PATH, name))]
	return jsonify({'sessions': sessions}), 200

# get/clear session history endpoint
@app.route('/api/history', methods=['GET', 'DELETE'])
def history_endpoint():
	session_id = request.headers.get('Session-ID')
	print("history request for session:", session_id)
	if not session_id:
		return jsonify({'error': 'Session-ID header is required'}), 400
	
	if request.method == 'DELETE':
		clear_session(session_id)
		# Clear session-specific chroma database
		clear_session_db(session_id)
		# also delete session folder
		session_folder = os.path.join(DATA_PATH, session_id)
		if os.path.exists(session_folder):
			shutil.rmtree(session_folder)
			print("Cleared history and database for session:", session_id)
			return jsonify({'message': 'Session cleared'}), 200
		else:
			return jsonify({'message': 'Session folder not found, but history and database cleared'}), 200
	
	if request.method == 'GET':
		'''
		export interface ChatMessage {
  id: string
  role: "user" | "assistant"
  content: string
  sources: string[]
  timestamp: Date
  isStreaming?: boolean
  truncatedContent?: string
}

		'''
		session_history = get_session_history(session_id)
		history = session_history.get_messages_with_timestamps()
		print("Fetched history for session:", session_id, history)
		history_serialized = [
			{
				'id': str(index),
				'role': 'assistant' if isinstance(msg, AIMessage) else 'user',
				'content': msg.content,
				'originalInput': session_history.get_message_metadata(index).get('originalInput'),
				'sources': session_history.get_message_metadata(index).get('sources', []),
				'timestamp': timestamp.isoformat(),
				'isStreaming': False,
				'truncatedContent': msg.content,
				'fileMetadata': session_history.get_message_metadata(index).get('fileMetadata'),
				'urls': session_history.get_message_metadata(index).get('urls'),
			}
			for (index, (msg, timestamp)) in enumerate(history)
		]
		return jsonify({'history': history_serialized}), 200

if __name__ == '__main__':
	app.run(port=PORT, debug=True)