"""
Document extraction module providing unified interface for extracting text from various file formats.

Supports:
- PDF files (text extraction + OCR for scanned pages)
- DOCX/DOC files (paragraphs, tables, headers, footers)
- Images (PNG, JPG, GIF, BMP, TIFF)
- Web URLs (screenshot + OCR or HTML parsing)
"""

import io
import os
import requests
import time
from typing import Optional, Tuple
from abc import ABC, abstractmethod
from PIL import Image
import numpy as np
from urllib.parse import urlparse
import logging
from langchain_community.document_loaders import WebBaseLoader
from selenium import webdriver
from selenium.webdriver.chrome.options import Options

# Optional imports with graceful degradation
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Configure logging
logger = logging.getLogger(__name__)

# Global OCR reader instance (initialized once and shared across all extractors)
_ocr_reader = None

def _initialize_global_ocr():
    """Initialize global EasyOCR reader."""
    global _ocr_reader
    if not EASYOCR_AVAILABLE or _ocr_reader is not None:
        return
    
    try:
        _ocr_reader = easyocr.Reader(['en'], gpu=False)
        print("Global EasyOCR reader initialized successfully")
    except Exception as e:
        print(f"Failed to initialize EasyOCR: {e}")
        _ocr_reader = None

def _get_ocr_reader():
    """Get the global OCR reader, initializing if necessary."""
    global _ocr_reader
    if _ocr_reader is None and EASYOCR_AVAILABLE:
        _initialize_global_ocr()
    return _ocr_reader


class DocumentExtractor(ABC):
    """Abstract base class for document text extractors."""
    
    @abstractmethod
    def can_extract(self, filename: str) -> bool:
        """Check if this extractor can handle the file."""
        pass
    
    @abstractmethod
    def extract(self, file_obj) -> str:
        """Extract text from file object. Returns annotated text."""
        pass


class ImageExtractor(DocumentExtractor):
    """Extracts text from images using EasyOCR."""
    
    SUPPORTED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff'}
    MIN_CONFIDENCE = 0.3
    
    def can_extract(self, filename: str) -> bool:
        """Check if file is a supported image format."""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        return ext in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_obj) -> str:
        """Extract text from image using OCR."""
        ocr_reader = _get_ocr_reader()
        if not ocr_reader:
            return f"[Image OCR unavailable - easyocr not installed]"
        
        try:
            filename = getattr(file_obj, 'filename', 'image')
            
            # Read image file into memory
            image_data = file_obj.read()
            file_obj.seek(0)
            
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Convert to numpy array for easyocr
            image_array = np.array(image)
            
            # Run OCR with confidence threshold
            ocr_result = ocr_reader.readtext(image_array, detail=1)
            
            if not ocr_result:
                return f"[Image File: {filename}]\nNo text detected in image."
            
            # Extract text from results, filter by confidence
            text_lines = [item[1] for item in ocr_result if item[2] > self.MIN_CONFIDENCE]
            text = '\n'.join(text_lines)
            
            if not text.strip():
                return f"[Image File: {filename}]\nNo text detected in image."
            
            logger.info(f"OCR extracted {len(text_lines)} lines from {filename}")
            return text
            
        except Exception as e:
            logger.error(f"Image OCR error for {filename}: {str(e)}")
            return f"[Image OCR Error: {str(e)}]"


class TextExtractor(DocumentExtractor):
    """Extracts text from plain text files."""
    
    SUPPORTED_EXTENSIONS = {'txt', 'text'}
    
    def can_extract(self, filename: str) -> bool:
        """Check if file is a plain text file."""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        return ext in self.SUPPORTED_EXTENSIONS
    
    def extract(self, file_obj) -> str:
        """Extract text from plain text file."""
        try:
            filename = getattr(file_obj, 'filename', 'document.txt')
            
            # Read file data
            file_data = file_obj.read()
            file_obj.seek(0)
            
            # Try to decode as text
            try:
                text = file_data.decode('utf-8')
            except UnicodeDecodeError:
                # Try alternative encodings
                for encoding in ['latin-1', 'utf-16', 'cp1252']:
                    try:
                        text = file_data.decode(encoding)
                        logger.info(f"Successfully decoded {filename} using {encoding}")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # If all encodings fail, return error
                    return f"[Text File: {filename}]\nCould not decode file. Please ensure it's a valid text file with UTF-8, Latin-1, UTF-16, or Windows-1252 encoding."
            
            if not text.strip():
                return f"[Text File: {filename}]\nFile is empty or contains only whitespace."
            
            logger.info(f"Extracted text from {filename}")
            return text
            
        except Exception as e:
            logger.error(f"Text extraction error for {filename}: {e}")
            return f"[Failed to extract text from file: {str(e)}]"


class PDFExtractor(DocumentExtractor):
    """Extracts text from PDF files using pypdf and OCR fallback."""
    
    MIN_CONFIDENCE = 0.3
    
    def can_extract(self, filename: str) -> bool:
        """Check if file is a PDF."""
        return filename.lower().endswith('.pdf')
    
    def extract(self, file_obj) -> str:
        """Extract text from PDF using text extraction + OCR fallback."""
        if not PDF_AVAILABLE:
            return "[PDF processing unavailable - pypdf not installed]"
        
        try:
            filename = getattr(file_obj, 'filename', 'document.pdf')
            pdf_reader = pypdf.PdfReader(file_obj)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    # Try text extraction first
                    page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                    else:
                        # If no text, try OCR fallback
                        text += self._try_page_ocr(file_obj, page_num, filename)
                        
                except Exception as e:
                    logger.error(f"Error extracting page {page_num + 1}: {e}")
                    text += f"[Failed to extract page {page_num + 1}: {str(e)}]\n"
            
            return text if text.strip() else "[PDF file had no extractable text]"
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return f"[Failed to process PDF: {str(e)}]"
    
    def _try_page_ocr(self, file_obj, page_num: int, filename: str) -> str:
        """Attempt OCR on a PDF page image."""
        ocr_reader = _get_ocr_reader()
        if not ocr_reader or not PYMUPDF_AVAILABLE:
            return f"[No text on page {page_num + 1}]\n"
        
        try:
            file_obj.seek(0)
            doc = fitz.open(stream=file_obj.read(), filetype="pdf")
            page_image = doc[page_num].get_pixmap(matrix=fitz.Matrix(2, 2))
            
            # Convert pixmap to numpy array
            image_data = page_image.tobytes("ppm")
            image = Image.open(io.BytesIO(image_data))
            image_array = np.array(image)
            
            # Run OCR on the page image
            ocr_result = ocr_reader.readtext(image_array, detail=1)
            ocr_text = '\n'.join([item[1] for item in ocr_result if item[2] > self.MIN_CONFIDENCE])
            
            result = f"\n--- Page {page_num + 1} (OCR) ---\n"
            result += ocr_text if ocr_text.strip() else "[No text detected on page]\n"
            
            doc.close()
            return result
            
        except Exception as e:
            logger.warning(f"OCR failed for page {page_num + 1}: {e}")
            return f"[OCR failed for page {page_num + 1}]\n"


class DOCXExtractor(DocumentExtractor):
    """Extracts text from DOCX/DOC files."""
    
    def can_extract(self, filename: str) -> bool:
        """Check if file is DOCX or DOC."""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        return ext in {'docx', 'doc'}
    
    def extract(self, file_obj) -> str:
        """Extract text from DOCX/DOC file."""
        filename = getattr(file_obj, 'filename', 'document')
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # Handle old .doc format
        if file_extension == 'doc':
            return f"[Old .doc format detected for {filename}. Please convert to .docx or PDF for better text extraction.]"
        
        if not DOCX_AVAILABLE:
            return "[DOCX processing unavailable - python-docx not installed]"
        
        try:
            # Read file into memory
            file_data = file_obj.read()
            file_obj.seek(0)
            
            # Load document from bytes
            doc = docx.Document(io.BytesIO(file_data))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Extract text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Header: {paragraph.text}]")
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Footer: {paragraph.text}]")
            
            # Combine all text
            text = '\n'.join(text_parts)
            
            if not text.strip():
                return f"[Document File: {filename}]\nNo text detected in document."
            
            logger.info(f"Extracted {len(text_parts)} text elements from {filename}")
            return text
            
        except Exception as e:
            logger.error(f"DOCX processing error for {filename}: {e}")
            return f"[Failed to process document: {str(e)}]"


class URLExtractor:
    """Intelligently extracts text from URLs using the best method."""
    
    TIMEOUT = 10
    MIN_CONTENT_LENGTH = 200  # Minimum chars to consider content "real"
    JS_INDICATORS = [
        # Meta frameworks
        'react', 'vue', 'angular', 'next.js', 'nuxt',
        # Common JS indicators in HTML
        '__NEXT_DATA__', '__NUXT__', 'ng-app', 'data-react-root',
        # Empty body indicators
        'id="root"', 'id="app"', 'id="__next"'
    ]
    
    def extract(self, url: str) -> Tuple[str, str]:
        """
        Extract text from URL, automatically choosing the best method.
        
        Returns:
            Tuple[str, str]: (extracted_text, method_used)
        """
        print(f"Processing URL: {url}")
        
        # Validate URL
        if not self._is_valid_url(url):
            return "[Invalid URL format]", "error"
        
        # Step 1: Quick static check
        needs_js = self._needs_javascript_rendering(url)
        
        if not needs_js:
            # Try static extraction first
            text, success = self._extract_static(url)
            if success:
                print(f"Static extraction succeeded ({len(text)} chars)")
                return text, "static"
            else:
                print("Static extraction failed, trying dynamic...")
        else:
            print("JS-heavy site detected, using dynamic extraction")
        
        # Step 2: Use Selenium if static failed or JS detected
        text, success = self._extract_dynamic(url)
        if success:
            print(f"Dynamic extraction succeeded ({len(text)} chars)")
            return text, "dynamic"
        
        return f"[Website: {url}]\nNo text content extracted.", "failed"
    
    def _is_valid_url(self, url: str) -> bool:
        """Validate URL format."""
        try:
            parsed = urlparse(url)
            return bool(parsed.scheme and parsed.netloc)
        except Exception:
            return False
    
    def _needs_javascript_rendering(self, url: str) -> bool:
        """
        Detect if a website likely needs JavaScript rendering.
        Uses heuristics based on HTML structure and meta tags.
        """
        try:
            # Get HTML without executing JS
            response = requests.get(
                url,
                timeout=self.TIMEOUT,
                headers={'User-Agent': 'Mozilla/5.0'}
            )
            response.raise_for_status()
            html = response.text.lower()
            
            # Check 1: Look for JS framework indicators
            for indicator in self.JS_INDICATORS:
                if indicator.lower() in html:
                    print(f"Found JS indicator: {indicator}")
                    return True
            
            # Check 2: Parse and analyze content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script/style tags
            for tag in soup(['script', 'style', 'noscript']):
                tag.decompose()
            
            # Get visible text
            body = soup.find('body')
            if not body:
                return True  # No body = likely JS-rendered
            
            text = body.get_text(strip=True)
            
            # Check 3: Too little content suggests JS rendering
            if len(text) < self.MIN_CONTENT_LENGTH:
                print(f"Minimal content detected ({len(text)} chars) - likely JS-rendered")
                return True
            
            # Check 4: Check for empty divs with common SPA IDs
            spa_roots = soup.find_all(id=['root', 'app', '__next'])
            if spa_roots:
                for root in spa_roots:
                    # If root div is empty or nearly empty
                    if not root.get_text(strip=True):
                        print("Found empty SPA root div")
                        return True
            
            print("Site appears to be static HTML")
            return False
            
        except Exception as e:
            logger.warning(f"Error detecting JS requirement: {e}")
            # Default to static if detection fails
            return False
    
    def _extract_static(self, url: str) -> Tuple[str, bool]:
        """Extract text using BeautifulSoup (fast, static HTML)."""
        try:
            print("Attempting static extraction with WebBaseLoader...")
            loader = WebBaseLoader(url)
            loader.requests_kwargs = {'timeout': self.TIMEOUT}
            docs = loader.load()
            
            if docs and docs[0].page_content:
                content = docs[0].page_content.strip()
                if len(content) >= self.MIN_CONTENT_LENGTH:
                    formatted = f"[Website Content from {url}]\n(Static HTML)\n\n{content}"
                    return formatted, True
            
            return "", False
            
        except Exception as e:
            logger.warning(f"Static extraction failed: {e}")
            return "", False
    
    def _extract_dynamic(self, url: str) -> Tuple[str, bool]:
        """Extract text using Selenium + OCR (slow, handles JS)."""
        try:
            # Get OCR reader
            ocr_reader = _get_ocr_reader()
            if not ocr_reader:
                print("OCR not available, trying basic Selenium text extraction")
                return self._extract_selenium_text_only(url)
            
            print("Attempting Selenium + OCR extraction...")
            
            # Setup Chrome
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--window-size=1920,1080')
            
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(url)
            time.sleep(3)  # Wait for JS to render
            
            # Try to get text content first
            body_text = driver.find_element("tag name", "body").text
            
            # If we got good text content, use that
            if len(body_text.strip()) >= self.MIN_CONTENT_LENGTH:
                driver.quit()
                formatted = f"[Website Content from {url}]\n(Rendered with JavaScript)\n\n{body_text}"
                return formatted, True
            
            # Otherwise, try OCR on screenshot
            screenshot_bytes = driver.get_screenshot_as_png()
            driver.quit()
            
            screenshot_image = Image.open(io.BytesIO(screenshot_bytes))
            if screenshot_image.mode != 'RGB':
                screenshot_image = screenshot_image.convert('RGB')
            
            image_array = np.array(screenshot_image)
            ocr_result = ocr_reader.readtext(image_array, detail=1)
            
            if ocr_result:
                text_lines = [item[1] for item in ocr_result if item[2] > 0.3]
                text = '\n'.join(text_lines)
                if text.strip():
                    formatted = f"[Website Content from {url}]\n(OCR from Screenshot)\n\n{text}"
                    return formatted, True
            
            return "", False
            
        except Exception as e:
            logger.warning(f"Dynamic extraction failed: {e}")
            return "", False
    
    def _extract_selenium_text_only(self, url: str) -> Tuple[str, bool]:
        """Fallback: Use Selenium without OCR."""
        try:
            print("Attempting Selenium text-only extraction...")
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(url)
            time.sleep(3)
            
            body_text = driver.find_element("tag name", "body").text
            driver.quit()
            
            if body_text.strip():
                formatted = f"[Website Content from {url}]\n(Rendered with JavaScript)\n\n{body_text}"
                return formatted, True
            
            return "", False
            
        except Exception as e:
            logger.warning(f"Selenium text extraction failed: {e}")
            return "", False


class DocumentExtractorFactory:
    """Factory for getting appropriate document extractor."""
    
    def __init__(self):
        self._extractors = [
            TextExtractor(),
            PDFExtractor(),
            DOCXExtractor(),
            ImageExtractor(),
        ]
    
    def extract_document(self, file_obj) -> str:
        """
        Extract text from document using appropriate extractor.
        
        Args:
            file_obj: File-like object with filename attribute
            
        Returns:
            Extracted text as string
        """
        filename = getattr(file_obj, 'filename', 'unknown')
        
        # Find matching extractor
        for extractor in self._extractors:
            if extractor.can_extract(filename):
                print(f"Using {extractor.__class__.__name__} for {filename}")
                return extractor.extract(file_obj)
        
        logger.warning(f"No extractor found for {filename}")
        file_extension = filename.lower().split('.')[-1] if '.' in filename else 'unknown'
        return f"[Unsupported file format: {file_extension}]"
    
    def extract_url(self, url: str) -> str:
        """
        Extract text from URL.
        
        Args:
            url: URL string to extract from
            
        Returns:
            Extracted text as string
        """
        extractor = URLExtractor()
        return extractor.extract(url)


# Global factory instance
_factory = DocumentExtractorFactory()


def extract_document(file_obj) -> str:
    """
    Extract text from document file.
    
    Args:
        file_obj: File-like object with filename attribute
        
    Returns:
        Extracted text as string with annotations
    """
    return _factory.extract_document(file_obj)


def extract_url(url: str) -> str:
    """
    Extract text from URL.
    
    Args:
        url: URL string
        
    Returns:
        Extracted text as string with annotations
    """
    return _factory.extract_url(url)
